import streamlit as st
from datetime import date
import io
import re
import os
import time 
import requests
import pandas as pd

# --- FILE PATHING & DIAGRAM MAPPING ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ASSETS_DIR = os.path.join(BASE_DIR, "diagrams")

# Static assets
AWS_PN_LOGO = os.path.join(ASSETS_DIR, "aws partner logo.jpg")
ONETURE_LOGO = os.path.join(ASSETS_DIR, "oneture logo1.jpg")
AWS_ADV_LOGO = os.path.join(ASSETS_DIR, "aws advanced logo1.jpg")

# Mapped Infra Costs
SOW_COST_TABLE_MAP = { 
    "L1 Support Bot POC SOW": { "poc_cost": "3,536.40 USD" }, 
    "Beauty Advisor POC SOW": { 
        "poc_cost": "4,525.66 USD + 200 USD (Amazon Bedrock Cost) = 4,725.66", 
        "prod_cost": "4,525.66 USD + 1,175.82 USD (Amazon Bedrock Cost) = 5,701.48" 
    }, 
    "Ready Search POC Scope of Work Document":{ "poc_cost": "2,641.40 USD" }, 
    "AI based Image Enhancement POC SOW": { "poc_cost": "2,814.34 USD" }, 
    "AI based Image Inspection POC SOW": { "poc_cost": "3,536.40 USD" }, 
    "Gen AI for SOP POC SOW": { "poc_cost": "2,110.30 USD" }, 
    "Project Scope Document": { "prod_cost": "2,993.60 USD" }, 
    "Gen AI Speech To Speech": { "prod_cost": "2,124.23 USD" }, 
    "PoC Scope Document": { "amazon_bedrock": "1,000 USD", "total": "$ 3,150" }
}

# AWS Calculator Links
CALCULATOR_LINKS = {
    "L1 Support Bot POC SOW": "https://calculator.aws/#/estimate?id=211ea64cba5a8f5dc09805f4ad1a1e598ef5238b",
    "Ready Search POC Scope of Work Document": "https://calculator.aws/#/estimate?id=f8bc48f1ae566b8ea1241994328978e7e86d3490",
    "AI based Image Enhancement POC SOW": "https://calculator.aws/#/estimate?id=9a3e593b92b796acecf31a78aec17d7eb957d1e5",
    "Beauty Advisor POC SOW": "https://calculator.aws/#/estimate?id=3f89756a35f7bac7b2cd88d95f3e9aba9be9b0eb",
    "AI based Image Inspection POC SOW": "https://calculator.aws/#/estimate?id=72c56f93b0c0e101d67a46af4f4fe9886eb93342",
    "Gen AI for SOP POC SOW": "https://calculator.aws/#/estimate?id=c21e9b242964724bf83556cfeee821473bb935d1",
    "Project Scope Document": "https://calculator.aws/#/estimate?id=37339d6e34c73596559fe09ca16a0ac2ec4c4252",
    "Gen AI Speech To Speech": "https://calculator.aws/#/estimate?id=8444ae26e6d61e5a43e8e743578caa17fd7f3e69",
    "PoC Scope Document": "https://calculator.aws/#/estimate?id=420ed9df095e7824a144cb6c0e9db9e7ec3c4153"
}

SOW_DIAGRAM_MAP = {
    "L1 Support Bot POC SOW": os.path.join(ASSETS_DIR, "L1 Support Bot POC SOW.png"),
    "Beauty Advisor POC SOW": os.path.join(ASSETS_DIR, "Beauty Advisor POC SOW.png"),
    "Ready Search POC Scope of Work Document": os.path.join(ASSETS_DIR, "Ready Search POC Scope of Work Document.png"),
    "AI based Image Enhancement POC SOW": os.path.join(ASSETS_DIR, "AI based Image Enhancement POC SOW.png"),
    "AI based Image Inspection POC SOW": os.path.join(ASSETS_DIR, "AI based Image Inspection POC SOW.png"),
    "Gen AI for SOP POC SOW": os.path.join(ASSETS_DIR, "Gen AI for SOP POC SOW.png"),
    "Project Scope Document": os.path.join(ASSETS_DIR, "Project Scope Document.png"),
    "Gen AI Speech To Speech": os.path.join(ASSETS_DIR, "Gen AI Speech To Speech.png"),
    "PoC Scope Document": os.path.join(ASSETS_DIR, "PoC Scope Document.png")
}

# --- CONFIGURATION ---
st.set_page_config(page_title="GenAI SOW Architect Pro", layout="wide", page_icon="📄")

st.markdown("""
    <style>
    .main { background-color: #f8fafc; }
    .stTabs [data-baseweb="tab"] { font-weight: 600; }
    .stakeholder-header { 
        background-color: #f1f5f9; padding: 8px 12px; border-radius: 6px; 
        margin-top: 10px; font-weight: bold; border-left: 4px solid #3b82f6;
    }
    .sow-preview {
        background-color: white; padding: 40px; border-radius: 12px;
        border: 1px solid #e2e8f0; line-height: 1.7; 
        color: #000000;
        font-family: "Times New Roman", Times, serif;
    }
    .sow-preview a { color: #0000EE; text-decoration: underline; }
    </style>
    """, unsafe_allow_html=True)

# Helper functions
def add_hyperlink(paragraph, text, url):
    from docx.oxml.shared import qn, OxmlElement
    import docx.opc.constants
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), '0000EE') 
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single')
    rPr.append(c); rPr.append(u); new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    new_run.append(t); hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def create_docx_logic(text_content, branding, sow_name):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    
    # PAGE 1: COVER
    p = doc.add_paragraph()
    if os.path.exists(AWS_PN_LOGO): doc.add_picture(AWS_PN_LOGO, width=Inches(1.6))
    doc.add_paragraph("\n" * 3)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(branding['sow_name']); run.font.size = Pt(26); run.bold = True; run.font.name = 'Times New Roman'
    stitle = doc.add_paragraph(); stitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stitle.add_run("Scope of Work Document").font.size = Pt(14)
    doc.add_paragraph("\n" * 4)
    
    l_table = doc.add_table(rows=1, cols=3); l_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if branding.get("customer_logo_bytes"):
        l_table.rows[0].cells[0].paragraphs[0].add_run().add_picture(io.BytesIO(branding["customer_logo_bytes"]), width=Inches(1.5))
    if os.path.exists(ONETURE_LOGO):
        l_table.rows[0].cells[1].paragraphs[0].add_run().add_picture(ONETURE_LOGO, width=Inches(1.8))
    if os.path.exists(AWS_ADV_LOGO):
        l_table.rows[0].cells[2].paragraphs[0].add_run().add_picture(AWS_ADV_LOGO, width=Inches(1.5))
    
    doc.add_paragraph("\n" * 3)
    dt = doc.add_paragraph(); dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt.add_run(branding["doc_date_str"]).bold = True
    doc.add_page_break()

    # Section Headers Mapping (Restructured Output to 5 Chapters)
    headers_map = {
        "1": "TABLE OF CONTENTS",
        "2": "PROJECT OVERVIEW",
        "3": "SCOPE OF WORK - TECHNICAL PROJECT PLAN",
        "4": "SOLUTION ARCHITECTURE / ARCHITECTURAL DIAGRAM",
        "5": "RESOURCES & COST ESTIMATES"
    }

    lines = text_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()
        if not line: i += 1; continue
        
        clean_line = re.sub(r'#+\s*', '', line).strip()
        upper = clean_line.upper()

        current_id = None
        for h_id, h_title in headers_map.items():
            if re.match(rf"^{h_id}[\.\s]+{re.escape(h_title)}", upper):
                current_id = h_id; break
        
        if current_id:
            if current_id != "1": doc.add_page_break()
            h = doc.add_heading(clean_line.upper(), level=1)
            for run in h.runs: 
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Injection Logic for Section 4 Architecture Diagram
            if current_id == "4":
                diag = SOW_DIAGRAM_MAP.get(sow_name)
                if diag and os.path.exists(diag):
                    doc.add_picture(diag, width=Inches(5.8))
                    p_cap = doc.add_paragraph(f"{sow_name} – Architecture Diagram")
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1; continue
            
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i]); i += 1
            if len(table_lines) >= 3:
                cols = [c.strip() for c in table_lines[0].split('|') if c.strip()]
                t_obj = doc.add_table(rows=1, cols=len(cols)); t_obj.style = "Table Grid"
                for idx, h_text in enumerate(cols):
                    cell = t_obj.rows[0].cells[idx]
                    r_h = cell.paragraphs[0].add_run(h_text)
                    r_h.bold = True; r_h.font.name = 'Times New Roman'
                for row_line in table_lines[2:]:
                    cells_data = [c.strip() for c in row_line.split('|') if c.strip()]
                    r_row = t_obj.add_row().cells
                    for idx, c_text in enumerate(cells_data):
                        if idx < len(r_row):
                            p_cell = r_row[idx].paragraphs[0]
                            # Insert Pricing link logic for Section 4
                            if "link" in c_text.lower() or "estimate" in c_text.lower():
                                add_hyperlink(p_cell, "AWS Calculator Link", CALCULATOR_LINKS.get(sow_name, "https://calculator.aws/"))
                            else:
                                p_cell.add_run(c_text).font.name = 'Times New Roman'
            continue

        if line.startswith('## '):
            h = doc.add_heading(clean_line, level=2)
            for run in h.runs: run.font.color.rgb = RGBColor(0, 0, 0); run.font.name = 'Times New Roman'
        elif line.startswith('### '):
            h = doc.add_heading(clean_line, level=3)
            for run in h.runs: run.font.color.rgb = RGBColor(0, 0, 0); run.font.name = 'Times New Roman'
        elif line.startswith('- ') or line.startswith('* '):
            p_b = doc.add_paragraph(style="List Bullet")
            p_b.add_run(re.sub(r'^[\-\*]\s*', '', line)).font.name = 'Times New Roman'
        else:
            p_n = doc.add_paragraph()
            p_n.add_run(clean_line).font.name = 'Times New Roman'
        i += 1
        
    bio = io.BytesIO(); doc.save(bio); return bio.getvalue()

def call_gemini_with_retry(payload, api_key_input=""):
    apiKey = api_key_input if api_key_input else ""
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={apiKey}"
    delays = [1, 2, 4]
    for attempt in range(len(delays)):
        try:
            res = requests.post(url, json=payload, timeout=30)
            if res.status_code == 200: return res, None
            if res.status_code in [503, 429]: time.sleep(delays[attempt]); continue
            return None, f"API Error {res.status_code}: {res.text}"
        except Exception: time.sleep(delays[attempt])
    return None, "Model overloaded."

# --- INITIALIZATION ---
def init_state():
    if 'generated_sow' not in st.session_state: st.session_state.generated_sow = ""
    if 'stakeholders' not in st.session_state:
        st.session_state.stakeholders = {
            "Partner": pd.DataFrame([{"Name": "Gaurav Kankaria", "Title": "Head of Analytics & ML", "Email": "gaurav.kankaria@oneture.com"}]),
            "Customer": pd.DataFrame([{"Name": "Prabhjot Singh", "Title": "Marketing Manager", "Email": "prabhjot.singh5@jublfood.com"}]),
            "AWS": pd.DataFrame([{"Name": "Anubhav Sood", "Title": "AWS Account Executive", "Email": "anbhsood@amazon.com"}]),
            "Escalation": pd.DataFrame([{"Name": "Omkar Dhavalikar", "Title": "AI/ML Lead", "Email": "omkar.dhavalikar@oneture.com"}])
        }
    if 'timeline_phases' not in st.session_state:
        st.session_state.timeline_phases = pd.DataFrame([
            {"Phase": "Infrastructure Setup", "Task": "Setup required AWS Services & gather documents", "Wk1": "X", "Wk2": "", "Wk3": "", "Wk4": ""},
            {"Phase": "Create Core Workflows", "Task": "Banner Upload & Validation / Compliance Flow", "Wk1": "X", "Wk2": "", "Wk3": "", "Wk4": ""},
            {"Phase": "Backend Components", "Task": "Build Compliance Engine & Tagging Module", "Wk1": "", "Wk2": "X", "Wk3": "X", "Wk4": "X"},
            {"Phase": "Feedback & Testing", "Task": "Test compliance accuracy against manual results", "Wk1": "", "Wk2": "", "Wk3": "", "Wk4": "X"}
        ])

init_state()

def reset_all():
    for key in list(st.session_state.keys()): del st.session_state[key]
    init_state(); st.rerun()

# --- INPUT PARAMETERS (STRICTLY PRESERVED) ---
with st.sidebar:
    st.image("https://img.icons8.com/fluency/96/artificial-intelligence.png", width=60)
    st.title("Architect Pro")
    api_key = st.text_input("Gemini API Key", type="password")
    st.divider()
    st.header("📋 1. Project Intake")
    sow_opts = list(SOW_DIAGRAM_MAP.keys())
    solution_type = st.selectbox("1.1 Solution Type", sow_opts)
    sow_key = solution_type
    engagement_type = st.selectbox("1.2 Engagement Type", ["PoC", "Pilot", "MVP", "Production"])
    industry_type = st.selectbox("1.3 Industry", ["Retail", "BFSI", "Manufacturing", "Telecom", "Healthcare", "Other"])
    final_industry = st.text_input("Specify Industry (if Other):") if industry_type == "Other" else industry_type
    if st.button("🗑️ Reset All", use_container_width=True): reset_all()

st.title("🚀 GenAI SOW Architect")
st.header("📸 Branding & Cover")
col_cov1, col_cov2 = st.columns(2)
with col_cov1: customer_logo = st.file_uploader("Upload Customer Logo", type=["png", "jpg", "jpeg"])
with col_cov2: doc_date = st.date_input("Document Date", date.today())

st.header("📄 2. Project Overview Section")
biz_objective = st.text_area("2.1 Business Objective", placeholder="Describe the business problem...", height=100)
sel_outcomes = st.multiselect("2.2 Key Outcomes Expected", ["Reduce effort", "Improve accuracy", "Cost reduction", "Compliance", "UX"])
st.subheader("2.3 Stakeholders Information")
st.session_state.stakeholders["Partner"] = st.data_editor(st.session_state.stakeholders["Partner"], use_container_width=True, key="ed_p")
st.session_state.stakeholders["Customer"] = st.data_editor(st.session_state.stakeholders["Customer"], use_container_width=True, key="ed_c")
st.session_state.stakeholders["AWS"] = st.data_editor(st.session_state.stakeholders["AWS"], use_container_width=True, key="ed_a")
st.session_state.stakeholders["Escalation"] = st.data_editor(st.session_state.stakeholders["Escalation"], use_container_width=True, key="ed_e")

st.header("📋 3. Assumptions & Dependencies")
sel_deps = st.multiselect("3.1 Customer Dependencies", ["Data availability", "API access", "SME availability", "AWS access"])
data_types = st.multiselect("3.2 Data Characteristics", ["Images", "Text", "Audio", "Video"])
key_assumptions = st.text_area("3.3 Key Assumptions", "PoC only, No production-grade SLA, Manual review for edge cases...")

st.header("🎯 4. Success Criteria")
sel_dims = st.multiselect("4.1 Success Dimensions", ["Accuracy", "Latency", "Cost efficiency"])
val_req = st.radio("4.2 User Validation Requirement", ["Required", "Not Required"])

st.header("🛠️ 5. Scope of Work – Functional Capabilities")
sel_caps = st.text_area("5.1 Core Capabilities", "Upload/Ingestion, Processing, Tagging Module...")
sel_ints = st.multiselect("5.2 Integrations Required", ["CRM", "ERP", "External APIs", "S3 Storage"])

st.header("🏢 6. Architecture & AWS Services")
compute_choices = st.selectbox("6.1 Compute & Orchestration", ["Lambda", "ECS", "Step Functions"])
ai_svcs = st.multiselect("6.2 GenAI / ML Services", ["Amazon Bedrock", "Textract", "Rekognition"])
st_svcs = st.multiselect("6.3 Storage & Search", ["S3", "OpenSearch", "DynamoDB"])
ui_layer = st.selectbox("6.4 UI Layer", ["Streamlit", "React", "Demo UI Only"])

st.header("⚙️ 7. Non-Functional Requirements")
perf = st.selectbox("7.1 Performance Expectations", ["Near real-time", "Batch", "SLA-based"])
sec = st.multiselect("7.2 Security & Compliance", ["IAM", "Encryption", "VPC", "SOC2 Alignment"])

st.header("📅 8. Timeline & Phasing")
poc_dur = st.selectbox("8.1 PoC Duration", ["2 weeks", "4 weeks", "6 weeks"])
st.session_state.timeline_phases = st.data_editor(st.session_state.timeline_phases, num_rows="dynamic", use_container_width=True, key="ed_t")

st.header("💰 9. Costing Inputs")
vol_input = st.text_input("9.1 Usage Volumes", "100 units/day")
ownership = st.selectbox("9.2 Cost Ownership", ["Jointly Funded by AWS & Oneture", "AWS Only", "Customer Only"])

st.header("🏁 10. Final Outputs")
delivs = st.multiselect("10.1 Deliverables", ["SOW Doc", "Demo UI", "Architecture Diagram", "Cost Estimate"])
nxt_steps = st.multiselect("10.2 Post-PoC Next Steps", ["Scale up", "Production proposal", "Refine Accuracy"])

# --- OUTPUT GENERATION LOGIC ---
if st.button("✨ Generate Professional SOW", type="primary", use_container_width=True):
    with st.spinner("Processing document sequence..."):
        def get_md(df): return df.to_markdown(index=False)
        cost_info = SOW_COST_TABLE_MAP.get(sow_key, {})
        cost_table = "| System | Infra Cost | AWS Cost Calculator Link |\n| --- | --- | --- |\n"
        for k, v in cost_info.items():
            cost_table += f"| {k.upper()} | {v} | link |\n"
        
        prompt = f"""
        Generate a professional enterprise SOW for {sow_key}. 
        STRICT MANDATE: Follow this 5-Chapter Output structure exactly:

        # 1 TABLE OF CONTENTS
        List the 5 sections below.

        # 2 PROJECT OVERVIEW
        ## 2.1 OBJECTIVE
        {biz_objective}
        ## 2.2 PROJECT SPONSOR(S) / STAKEHOLDER(S) / PROJECT TEAM
        ### Partner Executive Sponsor
        {get_md(st.session_state.stakeholders["Partner"])}
        ### Customer Executive Sponsor
        {get_md(st.session_state.stakeholders["Customer"])}
        ### AWS Executive Sponsor
        {get_md(st.session_state.stakeholders["AWS"])}
        ### Project Escalation Contacts
        {get_md(st.session_state.stakeholders["Escalation"])}
        ## 2.3 ASSUMPTIONS & DEPENDENCIES
        {key_assumptions}. Dependencies: {', '.join(sel_deps)}.
        ## 2.4 PROJECT SUCCESS CRITERIA
        Target dimensions: {', '.join(sel_dims)}. Validation: {val_req}.

        # 3 SCOPE OF WORK - TECHNICAL PROJECT PLAN
        {get_md(st.session_state.timeline_phases)}
        Note: POC would be demoed iteratively for all workflows across the {poc_dur}.

        # 4 SOLUTION ARCHITECTURE / ARCHITECTURAL DIAGRAM
        (Architecture Diagram Placement)
        ### Infrastructure Cost Table
        {cost_table}

        # 5 RESOURCES & COST ESTIMATES
        The project is {ownership}. Usage: {vol_input}. Outcomes: {', '.join(sel_outcomes)}.
        Deliverables: {', '.join(delivs)}. Next Steps: {', '.join(nxt_steps)}.
        """
        payload = {"contents": [{"parts": [{"text": prompt}]}]}
        res, err = call_gemini_with_retry(payload, api_key)
        if res:
            st.session_state.generated_sow = res.json()['candidates'][0]['content']['parts'][0]['text']
            st.rerun()
        else: st.error(err)

# --- PREVIEW SECTION ---
if st.session_state.generated_sow:
    st.divider()
    st.subheader("📄 Visual Preview")
    st.markdown('<div class="sow-preview">', unsafe_allow_html=True)
    
    p_content = st.session_state.generated_sow
    calc_url = CALCULATOR_LINKS.get(sow_key, "https://calculator.aws/")
    p_content = p_content.replace("link", f'[Estimate Link]({calc_url})')
    
    # Render Preview with Image Injection for Section 4
    if "# 4 SOLUTION ARCHITECTURE" in p_content:
        parts = p_content.split("# 4 SOLUTION ARCHITECTURE")
        st.markdown(parts[0] + "# 4 SOLUTION ARCHITECTURE", unsafe_allow_html=True)
        diag_out = SOW_DIAGRAM_MAP.get(sow_key)
        if diag_out and os.path.exists(diag_out):
            st.image(diag_out, caption=f"{sow_key} Architecture")
        st.markdown(parts[1], unsafe_allow_html=True)
    else:
        st.markdown(p_content, unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)
    
    if st.button("💾 Download Microsoft Word Document", use_container_width=True):
        branding = {"sow_name": sow_key, "customer_logo_bytes": customer_logo.getvalue() if customer_logo else None, "doc_date_str": doc_date.strftime("%d %B %Y")}
        docx_data = create_docx_logic(st.session_state.generated_sow, branding, sow_key)
        st.download_button("📥 Click to Download (.docx)", docx_data, f"SOW_{sow_key.replace(' ', '_')}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
